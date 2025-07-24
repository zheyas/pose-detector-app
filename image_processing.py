
import cv2
import mediapipe as mp
from PIL import Image
from docx import Document
from docx.shared import Inches

# Инициализация MediaPipe
mp_pose = mp.solutions.pose
pose = mp_pose.Pose(static_image_mode=True)
mp_drawing = mp.solutions.drawing_utils

# Словарь для перевода названий частей тела на русский язык
pose_translation = {
    "NOSE": "Нос",
    "LEFT_EYE_INNER": "Внутренний край левого глаза",
    "LEFT_EYE": "Левый глаз",
    "LEFT_EYE_OUTER": "Внешний край левого глаза",
    "RIGHT_EYE_INNER": "Внутренний край правого глаза",
    "RIGHT_EYE": "Правый глаз",
    "RIGHT_EYE_OUTER": "Внешний край правого глаза",
    "LEFT_EAR": "Левое ухо",
    "RIGHT_EAR": "Правое ухо",
    "MOUTH_LEFT": "Левая часть рта",
    "MOUTH_RIGHT": "Правая часть рта",
    "LEFT_SHOULDER": "Левое плечо",
    "RIGHT_SHOULDER": "Правое плечо",
    "LEFT_ELBOW": "Левый локоть",
    "RIGHT_ELBOW": "Правый локоть",
    "LEFT_WRIST": "Левое запястье",
    "RIGHT_WRIST": "Правое запястье",
    "LEFT_PINKY": "Левый мизинец",
    "RIGHT_PINKY": "Правый мизинец",
    "LEFT_INDEX": "Левый указательный палец",
    "RIGHT_INDEX": "Правый указательный палец",
    "LEFT_THUMB": "Левый большой палец",
    "RIGHT_THUMB": "Правый большой палец",

"LEFT_HIP": "Левый тазобедренный сустав",
"RIGHT_HIP": "Правый тазобедренный сустав",
"LEFT_KNEE": "Левое колено",
"RIGHT_KNEE": "Правое колено",
"LEFT_ANKLE": "Левая лодыжка",
"RIGHT_ANKLE": "Правая лодыжка",

    # добавьте остальные переводы
}

def process_image(image_path):
    image_bgr = cv2.imread(image_path)
    if image_bgr is None:
        print(f"Ошибка загрузки изображения: {image_path}")
        return None, None

    image_rgb = cv2.cvtColor(image_bgr, cv2.COLOR_BGR2RGB)
    results = pose.process(image_rgb)

    if not results.pose_landmarks:
        print("Не удалось обнаружить позы на изображении.")
        return None, None

    return image_bgr, results

def save_results_to_doc(image_bgr, file_name, landmarks):
    image_rgb = cv2.cvtColor(image_bgr, cv2.COLOR_BGR2RGB)
    pil_image = Image.fromarray(image_rgb)
    pil_image.save("annotated_image.png")

    doc = Document()
    doc.add_heading('Результаты анализа изображения', 0)
    doc.add_picture("annotated_image.png", width=Inches(5))

    doc.add_heading('Видимые части тела:', level=1)
    for idx, landmark in enumerate(landmarks.landmark):
        if landmark.visibility > 0.5:
            part_name = mp_pose.PoseLandmark(idx).name
            translated_name = pose_translation.get(part_name, part_name)
            doc.add_paragraph(f"{translated_name}")

    doc.save(file_name)

def describe_pose(landmarks):
    visible_parts = []
    for idx, landmark in enumerate(landmarks.landmark):
        if landmark.visibility > 0.5:
            part_name = mp_pose.PoseLandmark(idx).name
            translated_name = pose_translation.get(part_name, part_name)
            visible_parts.append(translated_name)

    if visible_parts:
        print("Видны части тела:", ", ".join(visible_parts))
    else:
        print("Не видно никаких частей тела.")
def draw_landmarks_on_image(image_bgr, pose_landmarks):
    mp_drawing.draw_landmarks(
        image_bgr,
        pose_landmarks,
        mp_pose.POSE_CONNECTIONS)
    return image_bgr


def is_face_fully_visible(landmarks, image_width, image_height):
    required_landmarks = [
        mp_pose.PoseLandmark.NOSE,
        mp_pose.PoseLandmark.LEFT_EYE,
        mp_pose.PoseLandmark.RIGHT_EYE,
        mp_pose.PoseLandmark.LEFT_EAR,
        mp_pose.PoseLandmark.RIGHT_EAR
    ]

    for landmark_id in required_landmarks:
        landmark = landmarks.landmark[landmark_id.value]

        # Проверяем видимость ключевых точек и их нахождение в центре изображения
        if landmark.visibility <= 0.5:
            return False
        if landmark.x < 0.1 or landmark.x > 0.9:  # Отступы от краев
            return False
        if landmark.y < 0.1 or landmark.y > 0.9:  # Отступы от краев
            return False

    return True